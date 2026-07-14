"""docx_reader.py — 读取 Word 文档内容"""
from __future__ import annotations

from skills import resolve_data_path


def docx_reader(path: str, data_root: str | None = None, max_chars: int = 5000) -> dict:
    """
    读取 Word（.docx）文档内容。

    优先用 python-docx，fallback 用 pandoc CLI 转 markdown，再 fallback 报错。
    path 相对于 data_root（即 data/ 目录）。

    返回 {"content": str, "num_paragraphs": int, "num_chars": int}
    """
    # 绝对路径直接用，不过 resolve_data_path
    from pathlib import Path as _Path
    _path_obj = _Path(path)
    if _path_obj.is_absolute():
        source = _path_obj
        root = _path_obj.parent
    else:
        source, root = resolve_data_path(path, data_root)

    if source.suffix.lower() not in {".docx", ".doc"}:
        raise ValueError(f"docx_reader only supports .docx/.doc files, got: {source.suffix}")
    if not source.is_file():
        raise FileNotFoundError(f"Word document not found: {path}")

    # ── 优先：python-docx ─────────────────────────────────────────────
    try:
        import docx  # python-docx
        document = docx.Document(str(source))
        paragraphs = [p.text for p in document.paragraphs]
        # 也提取表格中的文字
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
        content = "\n".join(p for p in paragraphs if p.strip())
        num_paragraphs = len([p for p in paragraphs if p.strip()])
        return {
            "content": content[:max_chars],
            "num_paragraphs": num_paragraphs,
            "num_chars": len(content[:max_chars]),
        }
    except ImportError:
        pass
    except Exception:
        pass

    # ── Fallback：pandoc CLI ──────────────────────────────────────────
    try:
        import subprocess
        import shutil
        if shutil.which("pandoc"):
            result = subprocess.run(
                ["pandoc", str(source), "-t", "markdown", "--wrap=none"],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0:
                content = result.stdout
                # 按段落粗算（空行分隔）
                paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
                num_paragraphs = len(paragraphs)
                return {
                    "content": content[:max_chars],
                    "num_paragraphs": num_paragraphs,
                    "num_chars": len(content[:max_chars]),
                }
            else:
                raise RuntimeError(f"pandoc 转换失败: {result.stderr[:200]}")
    except (ImportError, FileNotFoundError):
        pass
    except RuntimeError:
        raise
    except Exception:
        pass

    raise RuntimeError(
        f"无法读取 Word 文档 {path}：python-docx 和 pandoc 均不可用。"
        "请安装：pip install python-docx"
    )
