from __future__ import annotations

import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path


AGENT_ROOT = Path(__file__).resolve().parents[1]
CODE_ROOT = AGENT_ROOT / "code"
for entry in (str(AGENT_ROOT), str(CODE_ROOT)):
    if entry not in sys.path:
        sys.path.insert(0, entry)

from safe_python_exec import safe_python_exec
from skills.document_inspector import document_inspector


class RestrictedExecutorTests(unittest.TestCase):
    def test_success_returns_result_and_security_metadata(self):
        result = safe_python_exec("result = sum(range(10))", timeout_seconds=1)
        self.assertEqual(result["status"], "success")
        self.assertEqual(result["result"], 45)
        self.assertEqual(result["security"]["isolation"], "subprocess-json-stdio")

    def test_import_and_private_attribute_are_blocked(self):
        for source in (
            "import os",
            "result = (1).__class__",
            "result = open('secret.txt')",
            "result = getattr(1, '__class__')",
        ):
            with self.subTest(source=source):
                result = safe_python_exec(source, timeout_seconds=1)
                self.assertEqual(result["status"], "error")
                self.assertEqual(result["error"]["code"], "PERMISSION_DENIED")

    def test_memory_limit_terminates_worker(self):
        result = safe_python_exec(
            "result = [0] * 10000000",
            timeout_seconds=2,
            memory_limit_mb=32,
        )
        self.assertEqual(result["status"], "error")
        self.assertEqual(result["error"]["code"], "OVERFLOW")

    def test_timeout_terminates_worker(self):
        result = safe_python_exec("while True: pass", timeout_seconds=0.2)
        self.assertEqual(result["status"], "timeout")
        self.assertEqual(result["error"]["code"], "EXECUTION_TIMEOUT")

    def test_output_is_bounded(self):
        result = safe_python_exec(
            "print('x' * 1000)\nresult = 1",
            timeout_seconds=1,
            max_output_chars=100,
        )
        self.assertEqual(result["status"], "success")
        self.assertIn("output truncated", result["stdout"])
        self.assertLess(len(result["stdout"]), 150)

    def test_runtime_error_is_structured(self):
        result = safe_python_exec("result = 1 / 0", timeout_seconds=1)
        self.assertEqual(result["status"], "error")
        self.assertEqual(result["error"]["code"], "DIVISION_BY_ZERO")


class DocumentInspectorTests(unittest.TestCase):
    def test_text_chunks_include_line_evidence_and_matches(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            (root / "notes.md").write_text(
                "# Agent Tools\n\nA tool schema describes callable parameters.\n\nMemory is separate.",
                encoding="utf-8",
            )
            result = document_inspector(
                "notes.md",
                query="tool schema",
                chunk_size=100,
                chunk_overlap=10,
                data_root=str(root),
            )
        self.assertEqual(result["document"]["file_type"], "md")
        self.assertTrue(result["matches"])
        self.assertIn("line_start", result["matches"][0]["location"])
        self.assertEqual(len(result["document"]["sha256"]), 64)

    def test_csv_returns_table_and_row_evidence(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            (root / "scores.csv").write_text(
                "name,score\nalpha,91\nbeta,87\n",
                encoding="utf-8",
            )
            result = document_inspector(
                "scores.csv", query="alpha", data_root=str(root)
            )
        self.assertEqual(len(result["tables"]), 1)
        self.assertEqual(result["metadata"]["headers"], ["name", "score"])
        self.assertEqual(result["matches"][0]["location"], {"row": 2})

    def test_path_escape_is_rejected(self):
        with tempfile.TemporaryDirectory() as directory:
            with self.assertRaises(ValueError):
                document_inspector("../outside.txt", data_root=directory)

    @unittest.skipUnless(importlib.util.find_spec("docx"), "python-docx is unavailable")
    def test_docx_extracts_headings_and_tables(self):
        import docx

        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            path = root / "report.docx"
            document = docx.Document()
            document.add_heading("B2 Results", level=1)
            document.add_paragraph("The sandbox blocked unsafe imports.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "metric"
            table.cell(0, 1).text = "value"
            table.cell(1, 0).text = "blocked"
            table.cell(1, 1).text = "100%"
            document.save(path)
            result = document_inspector(
                "report.docx", query="unsafe imports", data_root=str(root)
            )
        self.assertEqual(result["metadata"]["headings"][0]["text"], "B2 Results")
        self.assertEqual(len(result["tables"]), 1)
        self.assertTrue(result["matches"])


if __name__ == "__main__":
    unittest.main(verbosity=2)
