"""Structured local document inspection with source-aware evidence chunks."""
from __future__ import annotations

import csv
import hashlib
import importlib.util
import re
from pathlib import Path
from typing import Any

from skills import resolve_data_path


_SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx", ".csv", ".tsv"}
_MAX_TABLES = 20
_MAX_TABLE_ROWS = 50
_MAX_TABLE_COLUMNS = 30


def _decode_text(data: bytes) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace"), "utf-8-replace"


def _clean_text(text: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).rstrip() for line in text.splitlines()]
    return "\n".join(lines).strip()


def _text_units(text: str) -> list[dict]:
    units: list[dict] = []
    current: list[str] = []
    start_line = 1
    lines = text.splitlines()
    for index, line in enumerate(lines, start=1):
        if line.strip():
            if not current:
                start_line = index
            current.append(line)
        elif current:
            units.append(
                {
                    "text": _clean_text("\n".join(current)),
                    "location": {"line_start": start_line, "line_end": index - 1},
                }
            )
            current = []
    if current:
        units.append(
            {
                "text": _clean_text("\n".join(current)),
                "location": {"line_start": start_line, "line_end": len(lines)},
            }
        )
    return [unit for unit in units if unit["text"]]


def _read_plain_text(source: Path) -> dict:
    text, encoding = _decode_text(source.read_bytes())
    return {
        "units": _text_units(text),
        "tables": [],
        "metadata": {"encoding": encoding},
        "num_pages": None,
        "warnings": [],
        "ocr": {"requested": False, "available": False, "pages_used": []},
    }


def _normalize_table(table: list[list[Any]], max_rows: int = _MAX_TABLE_ROWS) -> dict:
    rows = []
    for row in table[:max_rows]:
        values = [("" if cell is None else str(cell).strip()) for cell in row[:_MAX_TABLE_COLUMNS]]
        rows.append(values)
    return {
        "rows": rows,
        "num_rows_detected": len(table),
        "num_columns_detected": max((len(row) for row in table), default=0),
        "truncated": len(table) > max_rows or any(len(row) > _MAX_TABLE_COLUMNS for row in table),
    }


def _ocr_available() -> bool:
    return importlib.util.find_spec("pytesseract") is not None


def _ocr_pdf_page(page: Any, languages: str) -> str:
    import pytesseract

    image = page.to_image(resolution=150).original
    return pytesseract.image_to_string(image, lang=languages)


def _read_pdf(
    source: Path,
    extract_tables: bool,
    ocr_mode: str,
    ocr_languages: str,
) -> dict:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("PDF inspection requires pdfplumber") from exc

    units: list[dict] = []
    tables: list[dict] = []
    warnings: list[str] = []
    ocr_pages: list[int] = []
    available = _ocr_available()
    with pdfplumber.open(str(source)) as pdf:
        num_pages = len(pdf.pages)
        metadata = {str(key): value for key, value in (pdf.metadata or {}).items() if value is not None}
        for page_number, page in enumerate(pdf.pages, start=1):
            text = _clean_text(page.extract_text() or "")
            needs_ocr = ocr_mode == "force" or (ocr_mode == "auto" and len(text) < 20)
            if needs_ocr and available:
                try:
                    ocr_text = _clean_text(_ocr_pdf_page(page, ocr_languages))
                    if len(ocr_text) > len(text):
                        text = ocr_text
                        ocr_pages.append(page_number)
                except Exception as exc:
                    warnings.append(f"OCR failed on page {page_number}: {type(exc).__name__}: {exc}")
            elif needs_ocr and not available:
                warnings.append(f"page {page_number} has little text; optional OCR dependency is unavailable")
            if text:
                units.append({"text": text, "location": {"page": page_number}})

            if extract_tables and len(tables) < _MAX_TABLES:
                try:
                    for table_index, table in enumerate(page.extract_tables() or [], start=1):
                        if len(tables) >= _MAX_TABLES:
                            break
                        normalized = _normalize_table(table)
                        normalized.update({"page": page_number, "table_index": table_index})
                        tables.append(normalized)
                except Exception as exc:
                    warnings.append(
                        f"table extraction failed on page {page_number}: {type(exc).__name__}: {exc}"
                    )
    return {
        "units": units,
        "tables": tables,
        "metadata": metadata,
        "num_pages": num_pages,
        "warnings": warnings,
        "ocr": {"requested": ocr_mode != "off", "available": available, "pages_used": ocr_pages},
    }


def _read_docx(source: Path, extract_tables: bool) -> dict:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("DOCX inspection requires python-docx") from exc

    document = docx.Document(str(source))
    units: list[dict] = []
    headings: list[dict] = []
    current_section: str | None = None
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = _clean_text(paragraph.text)
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading"):
            current_section = text
            headings.append({"text": text, "paragraph": index, "style": style_name})
        location: dict[str, Any] = {"paragraph": index}
        if current_section:
            location["section"] = current_section
        units.append({"text": text, "location": location})

    tables: list[dict] = []
    if extract_tables:
        for table_index, table in enumerate(document.tables[:_MAX_TABLES], start=1):
            raw_table = [[cell.text for cell in row.cells] for row in table.rows]
            normalized = _normalize_table(raw_table)
            normalized["table_index"] = table_index
            tables.append(normalized)

    properties = document.core_properties
    metadata = {
        "title": properties.title or None,
        "author": properties.author or None,
        "subject": properties.subject or None,
        "headings": headings,
    }
    return {
        "units": units,
        "tables": tables,
        "metadata": metadata,
        "num_pages": None,
        "warnings": [],
        "ocr": {"requested": False, "available": False, "pages_used": []},
    }


def _read_delimited(source: Path, delimiter: str) -> dict:
    text, encoding = _decode_text(source.read_bytes())
    rows = list(csv.reader(text.splitlines(), delimiter=delimiter))
    normalized = _normalize_table(rows)
    table = {**normalized, "table_index": 1}
    headers = rows[0][:_MAX_TABLE_COLUMNS] if rows else []
    units = []
    for row_index, row in enumerate(rows[1 : _MAX_TABLE_ROWS + 1], start=2):
        values = [
            f"{headers[index] if index < len(headers) else f'column_{index + 1}'}: {value}"
            for index, value in enumerate(row[:_MAX_TABLE_COLUMNS])
        ]
        units.append({"text": " | ".join(values), "location": {"row": row_index}})
    return {
        "units": units,
        "tables": [table],
        "metadata": {"encoding": encoding, "delimiter": delimiter, "headers": headers},
        "num_pages": None,
        "warnings": [],
        "ocr": {"requested": False, "available": False, "pages_used": []},
    }


def _split_unit(text: str, chunk_size: int, chunk_overlap: int) -> list[tuple[int, int, str]]:
    if len(text) <= chunk_size:
        return [(0, len(text), text)]
    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        if end < len(text):
            boundary = max(text.rfind("\n", start, end), text.rfind("。", start, end), text.rfind(". ", start, end))
            if boundary > start + chunk_size // 2:
                end = boundary + 1
        chunks.append((start, end, text[start:end].strip()))
        if end >= len(text):
            break
        start = max(start + 1, end - chunk_overlap)
    return [chunk for chunk in chunks if chunk[2]]


def _build_chunks(
    units: list[dict],
    source_name: str,
    max_chars: int,
    chunk_size: int,
    chunk_overlap: int,
) -> tuple[list[dict], str, bool]:
    chunks: list[dict] = []
    content_parts: list[str] = []
    accepted_chars = 0
    truncated = False
    for unit_index, unit in enumerate(units, start=1):
        remaining = max_chars - accepted_chars
        if remaining <= 0:
            truncated = True
            break
        text = unit["text"]
        accepted = text[:remaining]
        if len(accepted) < len(text):
            truncated = True
        content_parts.append(accepted)
        for local_start, local_end, chunk_text in _split_unit(accepted, chunk_size, chunk_overlap):
            chunks.append(
                {
                    "chunk_id": f"chunk_{len(chunks) + 1:04d}",
                    "text": chunk_text,
                    "source": source_name,
                    "location": dict(unit["location"]),
                    "unit_index": unit_index,
                    "char_start": accepted_chars + local_start,
                    "char_end": accepted_chars + local_end,
                }
            )
        accepted_chars += len(accepted)
        if accepted_chars < max_chars:
            accepted_chars += 1
    return chunks, "\n".join(content_parts), truncated


def _query_terms(query: str) -> list[str]:
    lowered = query.lower()
    ascii_terms = re.findall(r"[a-z0-9_]+", lowered)
    cjk = re.findall(r"[\u4e00-\u9fff]", lowered)
    cjk_terms = cjk if len(cjk) < 2 else ["".join(cjk[index : index + 2]) for index in range(len(cjk) - 1)]
    return list(dict.fromkeys([*ascii_terms, *cjk_terms]))


def _rank_chunks(chunks: list[dict], query: str | None, top_k: int) -> list[dict]:
    if not query or not query.strip():
        return []
    terms = _query_terms(query)
    lowered_query = query.lower().strip()
    ranked = []
    for chunk in chunks:
        text = chunk["text"].lower()
        score = sum(text.count(term) for term in terms)
        if lowered_query and lowered_query in text:
            score += 5
        if score:
            ranked.append(
                {
                    "chunk_id": chunk["chunk_id"],
                    "score": score,
                    "text": chunk["text"],
                    "source": chunk["source"],
                    "location": chunk["location"],
                }
            )
    ranked.sort(key=lambda item: (-item["score"], item["chunk_id"]))
    return ranked[:top_k]


def document_inspector(
    path: str,
    query: str | None = None,
    max_chars: int = 20_000,
    chunk_size: int = 1000,
    chunk_overlap: int = 100,
    top_k: int = 5,
    extract_tables: bool = True,
    ocr_mode: str = "auto",
    ocr_languages: str = "chi_sim+eng",
    max_file_mb: int = 50,
    *,
    data_root: str | None = None,
) -> dict:
    """Inspect a local document and return structured, source-aware evidence."""
    if not isinstance(max_chars, int) or not 100 <= max_chars <= 200_000:
        raise ValueError("max_chars must be an integer between 100 and 200000")
    if not isinstance(chunk_size, int) or not 100 <= chunk_size <= 4000:
        raise ValueError("chunk_size must be an integer between 100 and 4000")
    if not isinstance(chunk_overlap, int) or not 0 <= chunk_overlap < chunk_size:
        raise ValueError("chunk_overlap must be non-negative and smaller than chunk_size")
    if not isinstance(top_k, int) or not 1 <= top_k <= 20:
        raise ValueError("top_k must be an integer between 1 and 20")
    if ocr_mode not in {"off", "auto", "force"}:
        raise ValueError("ocr_mode must be off, auto, or force")
    if not isinstance(max_file_mb, int) or not 1 <= max_file_mb <= 500:
        raise ValueError("max_file_mb must be an integer between 1 and 500")

    source, root = resolve_data_path(path, data_root)
    source = source.resolve()
    root = root.resolve()
    try:
        source.relative_to(root)
    except ValueError as exc:
        raise ValueError(f"path escapes data root: {path}") from exc
    if not source.is_file():
        raise FileNotFoundError(f"document not found: {path}")
    suffix = source.suffix.lower()
    if suffix not in _SUPPORTED_SUFFIXES:
        raise ValueError(f"unsupported document type: {suffix or '<none>'}")
    file_size = source.stat().st_size
    if file_size > max_file_mb * 1024 * 1024:
        raise ValueError(f"document exceeds {max_file_mb} MB limit")

    if suffix in {".txt", ".md"}:
        parsed = _read_plain_text(source)
    elif suffix == ".pdf":
        parsed = _read_pdf(source, bool(extract_tables), ocr_mode, ocr_languages)
    elif suffix == ".docx":
        parsed = _read_docx(source, bool(extract_tables))
    else:
        parsed = _read_delimited(source, "\t" if suffix == ".tsv" else ",")

    source_name = source.relative_to(root).as_posix()
    chunks, content, truncated = _build_chunks(
        parsed["units"], source_name, max_chars, chunk_size, chunk_overlap
    )
    matches = _rank_chunks(chunks, query, top_k)
    digest = hashlib.sha256(source.read_bytes()).hexdigest()
    title = parsed["metadata"].get("title") or source.stem
    return {
        "document": {
            "source": source_name,
            "file_type": suffix.lstrip("."),
            "title": title,
            "size_bytes": file_size,
            "sha256": digest,
            "num_pages": parsed["num_pages"],
            "num_units": len(parsed["units"]),
            "num_chunks": len(chunks),
            "num_chars": len(content),
            "truncated": truncated,
        },
        "metadata": parsed["metadata"],
        "content": content,
        "chunks": chunks,
        "tables": parsed["tables"],
        "matches": matches,
        "ocr": parsed["ocr"],
        "warnings": parsed["warnings"],
    }
